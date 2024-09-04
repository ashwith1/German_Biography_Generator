import os
import docx
import pandas as pd
from summarizer import Summarizer
import nltk
from nltk.tokenize import sent_tokenize

# Ensure the necessary nltk resource is available; download if not
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

def read_text_from_docx(file_path):
    """
    Read text from a Word document.
    """
    try:
        doc = docx.Document(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        return '\n'.join(paragraphs)
    except Exception as e:
        print(f"Error reading document: {e}")
        return None

def summarize_chunk(summarizer, chunk):
    """
    Generate a summary for a chunk of text.
    """
    return summarizer.generate_summary(chunk)

def generate_biography(summarizer, summary):
    """
    Generate a summary for a chunk of text.
    """
    return summarizer.generate_biography(summary)

def improve_coherence(summarizer, summary):
    """
    Improve the coherence of the biography.
    """
    return summarizer.improve_coherence(summary)

def divide_into_chunks(text, max_words_per_chunk):
    """
    split text into chunks close to specified maximum number of words without cutting through sentences.
    maintains context and integrity of content.
    If a sentence exceeds the word limit, it forms a chunk by itself.
    """
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = []
    current_word_count = 0

    for sentence in sentences: #splits entire text into individual sentences
        words_in_sentence = len(sentence.split())

        if current_word_count + words_in_sentence > max_words_per_chunk:
            chunks.append(' '.join(current_chunk))  #groups these sentences into chunks
            current_chunk = [sentence]  #create chunk if a particular sentence exceeds the word count
            current_word_count = words_in_sentence
        else:
            current_chunk.append(sentence)
            current_word_count += words_in_sentence

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

def remove_incomplete_sentence(biography):
    """
    Remove incomplete sentence at the end of the biography.
    """
    last_full_stop_index = biography.rfind('.')
    if last_full_stop_index != -1:
        return biography[:last_full_stop_index + 1]
    else:
        return biography

def read_text_from_csv(file_path):
    """
    Read text from a CSV file.
    """
    try:
        df = pd.read_csv(file_path, sep='\t')
        return df
    except Exception as e:
        print(f"Error reading CSV file: {e}")
        return None

def read_file(file_path):
    """
    Detect the file type and read the file accordingly.
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.docx':
        return read_text_from_docx(file_path)
    elif file_extension == '.csv':
        return read_text_from_csv(file_path)
    else:
        print("Unsupported file type")
        return None

def format_transkript_values(df):
    transkript_values = df['Transkript'].values
    formatted_text = "[{}]".format("\n".join(transkript_values))
    return formatted_text

def extract_rows_with_sprecher(df, sprecher_prefix):
    df = df.dropna(subset=['Sprecher'])
    filtered_rows = df[df['Sprecher'].str.startswith(sprecher_prefix)]
    transkript_list = filtered_rows['Transkript'].tolist()
    return transkript_list

# def extract_and_print_transcripts(csv_file_path, sprecher_prefix):
#     try:
#         df = pd.read_csv(csv_file_path, sep='\t')
#     except Exception as e:
#         print(f"Error reading CSV file: {e}")
#         return

#     if df is not None:
#         transkript_list = extract_rows_with_sprecher(df, sprecher_prefix)
#         if transkript_list:
#             print(f"Transkript data from rows with Sprecher starting with '{sprecher_prefix}':")
#             for transkript in transkript_list:
#                 print(transkript)
#         else:
#             print(f"No rows found with Sprecher starting with '{sprecher_prefix}'.")
#     else:
#         print("DataFrame is None. Cannot proceed.")

def biography(input_data, word_per_chunk=1000, output_file="summarized_document.docx"):
    if isinstance(input_data, pd.DataFrame):
        input_text = "\n".join(input_data['Transkript'].dropna().tolist())
    elif isinstance(input_data, str):
        input_text = input_data
    else:
        print("Unsupported input data type")
        return

    summarizer = Summarizer()
    biography_text = input_text
    while True:
        chunks = divide_into_chunks(biography_text, word_per_chunk)
        summary = ""
        for chunk in chunks:
            summary += summarize_chunk(summarizer, chunk) + "\n"
        total_tokens = len(summary.split())
        if total_tokens <= 600:
            break
        biography_text = summary

    summary = generate_biography(summarizer, summary)
    # summary = remove_incomplete_sentence(summary)
    print("This is summary : ", summary)

    # new_doc = docx.Document()
    # new_doc.add_paragraph(summary)
    # new_doc.save(output_file)

    return summary

if __name__ == "__main__":
    file_path = "C:\\Users\\asha4\\OneDrive - SRH\\Case Study-1\\Dennis- Files\\WG_ [EXTERN]  Transcripts and Biographies\\adg0001_er_2024_04_23.csv"
    sprecher_prefix = 'IP_'

# extract_and_print_transcripts(file_path, sprecher_prefix)


    df = read_file(file_path)
    if isinstance(df, pd.DataFrame):
        df = df.dropna(subset=['Sprecher', 'Transkript'])
        filtered_df = df[df['Sprecher'].str.startswith(sprecher_prefix)]
        biography(filtered_df)