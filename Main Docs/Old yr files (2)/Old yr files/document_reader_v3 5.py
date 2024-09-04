from summarizer import Summarizer
import docx


def read_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def summarize_chunk(summarizer, chunk):
    return summarizer.generate_summary(chunk)

def generate_biography(summarizer, summary):
    return summarizer.generate_biography(summary)


def divide_into_chunks(text, word_per_chunk):
    chunks = []
    words = text.split()
    num_chunks = len(words) // word_per_chunk
    for i in range(num_chunks):
        start = i * word_per_chunk
        end = (i + 1) * word_per_chunk
        chunks.append(" ".join(words[start:end]))
    # Adding the remaining words as the last chunk
    if len(words) % word_per_chunk != 0:
        chunks.append(" ".join(words[num_chunks * word_per_chunk:]))
    return chunks


def summarize_document(file_path, word_per_chunk=1000):
    summarizer = Summarizer()
    input_text = read_text_from_docx(file_path)
    chunks = divide_into_chunks(input_text, word_per_chunk)
    summarized_chunks = []
    for chunk in chunks:
        summary = summarize_chunk(summarizer, chunk)
        summarized_chunks.append(summary)
    # Combining the summaries into a new document
    new_doc = docx.Document()
    for summary in summarized_chunks:
        new_doc.add_paragraph(summary)
    new_doc.save("summarized_document.docx")
    # Checking the total number of tokens in the document
    total_tokens = sum(len(p.text.split()) for p in new_doc.paragraphs)

    if total_tokens > 600: #mistralai/Mixtral-8x7B-Instruct-v0.1, togethercomputer/StripedHyena-Nous-7B 
    # if total_tokens > 490: #snorkelai/Snorkel-Mistral-PairRM-DPO, mistralai/Mistral-7B-Instruct-v0.2
    # if total_tokens > 400: #upstage/SOLAR-10.7B-Instruct-v1.0
        # If total tokens exceed x, summarize it again
        summarize_document("summarized_document.docx", word_per_chunk)
    


if __name__ == "__main__":
    document_path = "Transkript_sample.docx"
    summarize_document(document_path)
