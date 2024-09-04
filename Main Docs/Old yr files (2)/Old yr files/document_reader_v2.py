# document_reader.py

from summarizer import Summarizer
import docx

def read_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def summarize_chunk(summarizer, chunk):
    summary = summarizer.generate_summary(chunk)
    return summary

def divide_into_chunks(text, chunk_size):
    chunks = []
    words = text.split()
    num_chunks = len(words) // chunk_size
    for i in range(num_chunks):
        start = i * chunk_size
        end = (i + 1) * chunk_size
        chunks.append(" ".join(words[start:end]))
    # Add the remaining words as the last chunk
    if len(words) % chunk_size != 0:
        chunks.append(" ".join(words[num_chunks * chunk_size:]))
    return chunks

def summarize_document(file_path, chunk_size=1000):
    summarizer = Summarizer()
    input_text = read_text_from_docx(file_path)
    chunks = divide_into_chunks(input_text, chunk_size)
    summarized_chunks = []
    for chunk in chunks:
        summary = summarize_chunk(summarizer, chunk)
        summarized_chunks.append(summary)
    # Combine the summaries into a new document
    new_doc = docx.Document()
    for summary in summarized_chunks:
        new_doc.add_paragraph(summary)
    new_doc.save("summarized_document.docx")

if __name__ == "__main__":
    # Example usage
    document_path = "Transkript_sample.docx"
    summarize_document(document_path)
